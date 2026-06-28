from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image as PILImage
from PIL import ImageDraw, ImageFont
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY, TA_LEFT
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import Image as RLImage
from reportlab.platypus import KeepTogether, Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "course_submission"
ASSET_DIR = OUT_DIR / "report_assets"
DOCX_PATH = OUT_DIR / "Minecraft_Constructing_Agents_课程报告.docx"
PDF_PATH = OUT_DIR / "Minecraft_Constructing_Agents_课程报告.pdf"
GITHUB_URL = "https://github.com/CityC196/Minecraft-Constructing-Agents"
SCREENSHOT_DIR = ROOT / "docs" / "assets" / "run-screenshots"


SCREENSHOT_CASES = [
    {
        "id": "A",
        "title": "现代两层家庭别墅",
        "prompt": "建一个现代两层家庭别墅，宽31深21，大玻璃窗，入口门厅，开放厨房，客厅，餐厅，三间卧室，书房，阳光房和彩色内饰，要有地毯、彩烛、盆栽、展示柜和清晰空间层次。",
    },
    {
        "id": "B",
        "title": "日式一层庭院住宅",
        "prompt": "建一个日式一层庭院住宅，宽29深23，木格栅，玄关，客厅，茶室，榻榻米卧室，小厨房，卫生间和枯山水庭院，要求内饰温暖缤纷，有灯笼、竹制家具、盆景和彩色地毯。",
    },
    {
        "id": "C",
        "title": "欧式三层大别墅",
        "prompt": "建一个欧式三层大别墅，宽39深29，对称侧翼，门廊，车库，客厅，餐厅，厨房，书房，四间卧室，阳台和卫生间，内饰豪华缤纷，有彩色地毯、旗帜、彩烛、盆栽和展示柜。",
    },
    {
        "id": "D",
        "title": "海滨架空度假住宅",
        "prompt": "建一个海滨架空度假住宅，宽33深21，抗风，抬高防潮，大露台，大玻璃客厅，厨房，餐厅，主卧，客房和书房，要求蓝白明亮内饰、户外座椅、信箱、无障碍坡道和遮雨平台。",
    },
]


PAPER = {
    "title": "Minecraft 建筑生成多智能体系统：从中文需求到可执行数据包",
    "authors": "龙想 2300011196；石宇宸 2300011051",
    "course": "《大语言模型与信息决策》课程项目",
    "date": "2026 年 6 月",
    "abstract": (
        "Minecraft 建筑生成任务同时涉及自然语言理解、空间规划、材料选择、几何约束、"
        "连通性校验和游戏数据包导出。若直接要求大语言模型输出具体方块坐标，系统容易出现"
        "非法材料、坐标不一致、房间不连通和命令不可执行等问题。本文介绍一套面向 Minecraft "
        "Java 的建筑生成多智能体系统。系统将大语言模型限制在高层语义决策层，使其输出风格、"
        "体块、材料、房间拓扑和设计意图等 JSON；随后由本地确定性几何引擎完成 CSG 体块生成、"
        "BSP 房间划分、A* 连通路径、室内装饰、质量校验、自动修复和数据包导出。项目从最初的"
        "单一建筑智能体逐步迭代为多智能体流水线，并逐步加入语义规划、本地几何、质量检查、"
        "可视化预览和候选择优机制。当前版本支持无 API key 的 mock 模式，可生成 blueprint.json、"
        "architect_datapack、preview.html 和 run_report.md；全量测试结果为 176 项通过、0 项失败。"
    ),
    "keywords": "大语言模型智能体；Minecraft；多智能体系统；CSG；BSP；A*；数据包",
    "sections": [
        {
            "heading": "1 引言",
            "paragraphs": [
                "Minecraft 的开放世界为生成式建筑提供了直观的实验场景。用户可以用一句中文描述建筑需求，例如“建一个欧式大房子”或“建造一个木屋别墅”，但系统真正需要解决的问题远不止文本生成。一个可用的建筑结果必须包含合理尺度、入口、房间、楼梯、屋顶、窗、材料、内饰、路径和最终可执行的 Minecraft 命令。",
                "项目早期的设想是让智能体直接把需求变成建筑蓝图或命令。实际尝试后，我们发现这一路径不稳定：大模型擅长描述建筑意图，却不擅长长期保持三维坐标、材料注册名和拓扑约束的一致性。由此，项目的核心问题被重新定义为：如何让大模型负责语义与设计决策，同时让本地程序负责确定性几何落地和工程校验。",
                "本文围绕这一修改后的问题定义展开。系统目标不是制作一个实时连服机器人，而是构建从中文需求到 Minecraft 可安装数据包的完整、可复现、可测试流水线。"
            ],
        },
        {
            "heading": "2 系统总体设计",
            "paragraphs": [
                "系统采用“LLM 语义智能体与本地确定性几何引擎”相结合的混合架构，如图 1 所示。大模型主要负责开放需求理解、风格构思和高层 JSON 决策；本地程序负责坐标、材料、连通性、命令体积和数据包导出等确定性任务。",
                "这种边界划分是项目最重要的构思修改：我们没有让模型直接写坐标，而是让它提出建筑意图，再由程序把意图变成可检查的 Minecraft 方块网格。"
            ],
            "figures": [
                {
                    "id": "overall_architecture",
                    "caption": "图 1 系统总体架构：LLM 负责语义决策，本地程序负责几何落地与导出",
                }
            ],
            "table": {
                "caption": "表 1 系统层次与模块职责",
                "headers": ["层次", "代表模块", "主要职责"],
                "widths": [1.15, 2.1, 3.25],
                "rows": [
                    ["语义层", "ArchitectAgent, PlannerAgent", "将中文需求转化为风格、材料、体块和房间拓扑 JSON"],
                    ["几何层", "CSGBuilder, BSPPartitioner, AStarPathfinder", "生成外壳、房间、门洞、楼梯和可达路径"],
                    ["质量层", "QA, Repair, Optimizer", "校验材料、连通性、入口、装饰和命令数量"],
                    ["导出层", "Exporter", "输出 blueprint.json、mcfunction 和可安装 datapack"],
                ],
            },
        },
        {
            "heading": "3 方法",
            "paragraphs": [
                "图 2 展示了主流程。流程分为语义决策、确定性几何、质量闭环和数据包导出四段；每段都有明确输入输出，因此可以单独测试和回退。"
            ],
            "subsections": [
                {
                    "heading": "3.1 语义智能体分工",
                    "paragraphs": [
                        "主流程位于 construction_method_v1。用户输入首先被解析为 buildSpec，随后 ArchitectAgent 生成外壳语义，MaterialPaletteAgent 基于 Minecraft Java 1.21.1 方块目录校验材料，PlannerAgent 生成房间节点与连通边，CreativeDesignAgent 决定体块变体、平面排序、屋顶、立面和场地策略。",
                        "后续智能体继续细化建筑表现。StructureAgent 给出支撑和荷载路径语义，FacadeAgent、RoofAgent、SiteLandscapeAgent 分别规划立面、屋顶和场地，InteriorDetailAgent 与 DecoratorAgent 负责房间级家具、灯光、植物和功能台面。整个过程中，LLM 的输出被约束为 JSON 语义对象，而不是直接写入坐标命令。"
                    ],
                },
                {
                    "heading": "3.2 确定性几何落地",
                    "paragraphs": [
                        "几何层是系统稳定性的关键。CSGBuilder 将主体、侧翼、门廊、塔楼等语义体块合并为稀疏 voxel 网格，并抽取外表面、楼板、屋顶和开窗。BSPPartitioner 根据楼层功能和房间权重划分室内空间；AStarPathfinder 根据拓扑关系生成门洞、楼梯和可通行路径。这样可以避免 LLM 直接坐标生成中的穿墙、断路和房间重叠问题。",
                        "导出层将最终网格转换为 Minecraft 数据包。玩家复制 architect_datapack 到世界 datapacks 目录后，执行 /reload 和 /function architect:run 即可触发 clear + build。系统同时输出 preview.html 和 run_report.md，便于不进入游戏时检查结构、流程和统计信息。"
                    ],
                    "figures": [
                        {
                            "id": "generation_pipeline",
                            "caption": "图 2 从中文 prompt 到 Minecraft datapack 的主流程",
                        }
                    ],
                },
                {
                    "heading": "3.3 质量检查与可复现导出",
                    "paragraphs": [
                        "在生成末端，系统不会直接信任任何单个智能体输出，而是把结构、材料、入口、房间可达性、装饰和命令数量放入统一检查。若发现缺入口、断路、非法方块或命令过多，Repair 与 Optimizer 会在导出前进行修复和压缩。",
                        "如图 3 所示，质量闭环的目标是让一次中文 prompt 最终落到可复现的工程产物：同一个 seed 可以复现同一份 blueprint、预览页、运行报告和 Minecraft 数据包。"
                    ],
                    "figures": [
                        {
                            "id": "quality_loop",
                            "caption": "图 3 质量检查与可复现导出闭环",
                        }
                    ],
                },
            ],
        },
        {
            "heading": "4 实现过程与构思迭代",
            "paragraphs": [
                "项目的主要工作并不是一次性写出最终架构，而是在多轮实现中逐步修改问题定义。图 4 概括了 Git 历史中几个关键阶段：从初始 agent，到多智能体拆分、语义规划、本地几何、质量检查和提交版本整理。",
                "这些变化反映了项目构思的转向。最初我们关注“如何生成一个房子”，后来逐步意识到课程项目更需要展示智能体系统如何分工、如何约束大模型、如何把开放语义落地为可靠工程产物。"
            ],
            "figures": [
                {
                    "id": "iteration_timeline",
                    "caption": "图 4 项目构思与实现迭代路线",
                }
            ],
        },
        {
            "heading": "5 实验与结果",
            "paragraphs": [
                "当前全量测试为 176 项通过、0 项失败，其中包括课程提交文档测试；原项目功能基线为 173 项通过、0 项失败。测试覆盖 Architect fallback、Planner、CSG、BSP、A*、Decorator、候选择优、LLM provider、材料目录和可视化输出等模块。",
                "本地样例 out/2026-06-19-145532212 使用 prompt“建造一个木屋别墅”。run_report.md 记录了 LLM 调用状态、seed、几何统计、装饰数量、QA 状态和 Minecraft 使用步骤。该样例说明系统并非只生成文本，而是在输出前经过结构、材料、连通和命令层面的检查。",
                "为展示实际运行效果，我们补充了四组 Minecraft 游戏内截图。A-D 分别对应现代两层家庭别墅、日式一层庭院住宅、欧式三层大别墅和海滨架空度假住宅；每组 1、2 为外景，3 为内饰。"
            ],
            "table": {
                "caption": "表 2 本地样例运行结果",
                "headers": ["指标", "结果", "说明"],
                "widths": [1.5, 1.65, 3.35],
                "rows": [
                    ["建筑尺寸", "43 x 16 x 57", "来自 run_report.md 的几何校验结果"],
                    ["房间可达性", "14/14", "入口可达房间全部通过"],
                    ["QA 检查", "9/9", "结构、材料、入口和装饰检查通过"],
                    ["装饰数量", "819", "由室内与风格装饰模块写入"],
                    ["命令压缩", "2692 -> 1251", "压缩倍率约 2.15x，降低数据包函数体积"],
                    ["建造入口", "/function architect:run", "数据包内置清理与建造函数"],
                ],
            },
            "figures_after_table": [
                {
                    "id": "run_exterior_grid",
                    "caption": "图 5 四个 prompt 的外景运行截图（A-D，每组两张外景）",
                },
                {
                    "id": "run_interior_grid",
                    "caption": "图 6 四个 prompt 的室内运行截图（A-D，每组一张内饰）",
                },
            ],
        },
        {
            "heading": "6 讨论",
            "paragraphs": [
                "从课程要求看，项目的整体性体现在所有模块都围绕同一条主线协作：中文需求输入、智能体语义决策、本地几何落地、质量检查和数据包导出。项目的创新性主要体现在三点。第一，系统清晰区分 LLM 与程序算法的职责，避免让大模型直接处理坐标和命令。第二，语义 JSON 与本地几何之间有清晰接口，便于单独测试和回退。第三，项目保留 mock 兜底和自动化测试，使没有 token 或现场网络时仍能复现完整流程。",
                "项目也存在局限。当前版本不是 Mineflayer 实时机器人，也不模拟玩家逐块放置；GDMC HTTP 客户端虽然保留，但主交付仍是数据包导出。自动检查可以发现结构、材料和连通问题，却不能完全替代玩家视角下的尺度、镜头、路径和 block readability 检查。后续可补充真实游戏截图、人工评分样本，以及更严格的视觉检查。"
            ],
        },
        {
            "heading": "7 结论",
            "paragraphs": [
                "Minecraft Constructing Agents 构建了一条从中文建筑需求到 Minecraft 可执行数据包的完整流水线。项目的核心经验是：在开放生成任务中，大语言模型更适合作为语义和设计决策者，而不是直接坐标生成器；本地几何算法、材料目录、连通性校验和命令优化可以为其提供稳定落地能力。",
                "通过多智能体分工、CSG/BSP/A* 几何引擎和质量闭环，项目将一个模糊的建房需求转化为可复现、可检查、可安装的工程产物。相比最终单次生成效果，项目更重要的贡献在于展示了一个课程级智能体系统如何从想法、失败尝试和架构调整中逐步形成。"
            ],
        },
        {
            "heading": "AI 辅助说明",
            "paragraphs": [
                "本项目合理使用 AI 工具辅助开发和文档整理。AI 辅助主要用于代码修改建议、文档初稿整理、报告润色、测试清单和网页排版检查；项目方向、系统边界、架构取舍、运行调试、测试执行、结果确认和最终整合由小组完成。报告和网站不伪造运行截图；若最终版本需要图像材料，应使用 preview.html 或 Minecraft 游戏内真实运行结果截图。"
            ],
        },
        {
            "heading": "参考资料",
            "references": [
                "Minecraft-Constructing-Agents 项目仓库与 README，https://github.com/CityC196/Minecraft-Constructing-Agents。",
                "本地样例运行报告：out/2026-06-19-145532212/run_report.md。",
                "课程提交展示页与提交清单：docs/index.html，SUBMISSION.md。",
                "推荐 Prompt 库与截图建议：docs/recommended-prompts.md。"
            ],
        },
    ],
}


def image_path(figure_id):
    return ASSET_DIR / f"{figure_id}.png"


def load_image_font(size, bold=False):
    candidates = [
        Path("C:/Windows/Fonts/msyhbd.ttc") if bold else Path("C:/Windows/Fonts/msyh.ttc"),
        Path("C:/Windows/Fonts/simhei.ttf") if bold else Path("C:/Windows/Fonts/Deng.ttf"),
        Path("C:/Windows/Fonts/simsun.ttc"),
    ]
    for path in candidates:
        if path.exists():
            return ImageFont.truetype(str(path), size=size)
    return ImageFont.load_default()


def text_size(draw, text, font):
    bbox = draw.textbbox((0, 0), text, font=font)
    return bbox[2] - bbox[0], bbox[3] - bbox[1]


def wrap_text(draw, text, font, max_width):
    lines = []
    for raw_line in text.split("\n"):
        current = ""
        for char in raw_line:
            candidate = current + char
            if text_size(draw, candidate, font)[0] <= max_width or not current:
                current = candidate
            else:
                lines.append(current)
                current = char
        if current:
            lines.append(current)
    return lines


def draw_centered_text(draw, box, text, font, fill="#17212F", line_gap=8):
    x1, y1, x2, y2 = box
    lines = wrap_text(draw, text, font, x2 - x1 - 34)
    heights = [text_size(draw, line, font)[1] for line in lines]
    total = sum(heights) + line_gap * (len(lines) - 1)
    y = y1 + (y2 - y1 - total) / 2
    for line, height in zip(lines, heights):
        width = text_size(draw, line, font)[0]
        draw.text((x1 + (x2 - x1 - width) / 2, y), line, font=font, fill=fill)
        y += height + line_gap


def draw_round_box(draw, box, text, font, fill, outline="#8AA1B7", text_fill="#17212F", radius=18):
    draw.rounded_rectangle(box, radius=radius, fill=fill, outline=outline, width=2)
    draw_centered_text(draw, box, text, font, fill=text_fill)


def draw_arrow(draw, start, end, fill="#345B7D", width=4):
    x1, y1 = start
    x2, y2 = end
    draw.line((x1, y1, x2, y2), fill=fill, width=width)
    if abs(x2 - x1) >= abs(y2 - y1):
        direction = 1 if x2 >= x1 else -1
        head = [(x2, y2), (x2 - direction * 18, y2 - 10), (x2 - direction * 18, y2 + 10)]
    else:
        direction = 1 if y2 >= y1 else -1
        head = [(x2, y2), (x2 - 10, y2 - direction * 18), (x2 + 10, y2 - direction * 18)]
    draw.polygon(head, fill=fill)


def new_canvas(width=1800, height=860):
    image = PILImage.new("RGB", (width, height), "#FFFFFF")
    draw = ImageDraw.Draw(image)
    return image, draw


def save_diagram(image, figure_id):
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    image.save(image_path(figure_id), "PNG")


def cover_resize(image, size):
    target_w, target_h = size
    source_w, source_h = image.size
    scale = max(target_w / source_w, target_h / source_h)
    resized = image.resize((round(source_w * scale), round(source_h * scale)), PILImage.Resampling.LANCZOS)
    left = (resized.width - target_w) // 2
    top = (resized.height - target_h) // 2
    return resized.crop((left, top, left + target_w, top + target_h))


def draw_screenshot_card(canvas, draw, screenshot, box, label, label_font, label_fill="#17212F"):
    x1, y1, x2, y2 = box
    label_h = 46
    draw.rounded_rectangle((x1, y1, x2, y2), radius=12, fill="#FFFFFF", outline="#B8C3CC", width=2)
    image_x1 = x1 + 12
    image_y1 = y1 + 12
    image_x2 = image_x1 + screenshot.width
    image_y2 = image_y1 + screenshot.height
    canvas.paste(screenshot, (image_x1, image_y1))
    draw.rectangle((image_x1, image_y1, image_x2, image_y2), outline="#B8C3CC", width=1)
    draw.text((x1 + 16, y2 - 34), label, font=label_font, fill=label_fill)


def make_screenshot_grid(figure_id, image_items, title, subtitle, columns=2):
    thumb_size = (760, 475)
    card_w = thumb_size[0] + 24
    card_h = thumb_size[1] + 68
    gap_x = 34
    gap_y = 28
    margin_x = 70
    top = 145
    rows = (len(image_items) + columns - 1) // columns
    width = margin_x * 2 + columns * card_w + (columns - 1) * gap_x
    height = top + rows * card_h + (rows - 1) * gap_y + 60
    image, draw = new_canvas(width=width, height=height)
    title_font = load_image_font(34, bold=True)
    sub_font = load_image_font(20)
    label_font = load_image_font(20, bold=True)
    draw.text((70, 46), title, font=title_font, fill="#17212F")
    draw.text((72, 94), subtitle, font=sub_font, fill="#566170")

    for idx, item in enumerate(image_items):
        row = idx // columns
        col = idx % columns
        x1 = margin_x + col * (card_w + gap_x)
        y1 = top + row * (card_h + gap_y)
        src = SCREENSHOT_DIR / f"{item['image']}.jpg"
        with PILImage.open(src) as screenshot:
            screenshot = cover_resize(screenshot.convert("RGB"), thumb_size)
        draw_screenshot_card(
            image,
            draw,
            screenshot,
            (x1, y1, x1 + card_w, y1 + card_h),
            item["label"],
            label_font,
        )
    save_diagram(image, figure_id)


def draw_title(draw, title, subtitle=None):
    title_font = load_image_font(36, bold=True)
    sub_font = load_image_font(22)
    draw.text((70, 46), title, font=title_font, fill="#17212F")
    if subtitle:
        draw.text((72, 96), subtitle, font=sub_font, fill="#566170")


def make_overall_architecture():
    image, draw = new_canvas()
    title_font = load_image_font(30, bold=True)
    box_font = load_image_font(24, bold=True)
    note_font = load_image_font(20)
    draw_title(draw, "总体架构", "让 LLM 做语义决策，让本地程序做可验证落地")

    boxes = [
        ((90, 230, 360, 390), "中文需求\n自然语言 prompt", "#F6F8FA"),
        ((470, 185, 790, 435), "LLM 语义智能体\nArchitect / Planner\nCreativeDesign\n输出高层 JSON", "#EAF3FF"),
        ((900, 185, 1220, 435), "本地几何引擎\nCSG / BSP / A*\n坐标、房间、路径", "#EAF7EF"),
        ((1330, 185, 1650, 435), "质量与导出\nQA / Repair / Optimizer\nMinecraft datapack", "#FFF4DF"),
    ]
    for box, text, fill in boxes:
        draw_round_box(draw, box, text, box_font, fill)
    for start, end in [((360, 310), (470, 310)), ((790, 310), (900, 310)), ((1220, 310), (1330, 310))]:
        draw_arrow(draw, start, end)

    layers = [
        ((470, 540, 790, 640), "开放语义：风格、体块、材料、拓扑", "#EAF3FF"),
        ((900, 540, 1220, 640), "确定性约束：坐标、注册名、连通性", "#EAF7EF"),
        ((1330, 540, 1650, 640), "可提交结果：blueprint + mcfunction + datapack", "#FFF4DF"),
    ]
    for box, text, fill in layers:
        draw_round_box(draw, box, text, note_font, fill, outline="#B8C3CC", radius=14)
    draw.text((90, 575), "核心边界", font=title_font, fill="#1F3A5F")
    save_diagram(image, "overall_architecture")


def make_generation_pipeline():
    image, draw = new_canvas(height=820)
    box_font = load_image_font(22, bold=True)
    small_font = load_image_font(18)
    draw_title(draw, "主生成流程", "每一步都有明确输入输出，便于测试、回退和复现")

    top_y = 165
    boxes = [
        ((70, top_y, 260, top_y + 110), "用户 prompt", "#F6F8FA"),
        ((330, top_y, 520, top_y + 110), "buildSpec\n尺寸与偏好", "#F6F8FA"),
        ((590, top_y, 810, top_y + 110), "Architect\n外壳语义", "#EAF3FF"),
        ((880, top_y, 1100, top_y + 110), "Planner\n房间拓扑", "#EAF3FF"),
        ((1170, top_y, 1390, top_y + 110), "CreativeDesign\n体块/屋顶/场地", "#EAF3FF"),
        ((1460, top_y, 1700, top_y + 110), "语义 JSON", "#EAF3FF"),
    ]
    for box, text, fill in boxes:
        draw_round_box(draw, box, text, box_font, fill)
    for idx in range(len(boxes) - 1):
        draw_arrow(draw, (boxes[idx][0][2], top_y + 55), (boxes[idx + 1][0][0], top_y + 55))

    mid_y = 365
    boxes2 = [
        ((160, mid_y, 430, mid_y + 120), "CSGBuilder\n外壳与屋顶", "#EAF7EF"),
        ((535, mid_y, 805, mid_y + 120), "BSPPartitioner\n房间划分", "#EAF7EF"),
        ((910, mid_y, 1180, mid_y + 120), "AStarPathfinder\n门洞与楼梯", "#EAF7EF"),
        ((1285, mid_y, 1555, mid_y + 120), "Decorator\n家具与细节", "#EAF7EF"),
    ]
    for box, text, fill in boxes2:
        draw_round_box(draw, box, text, box_font, fill)
    draw_arrow(draw, (1580, top_y + 110), (295, mid_y))
    for idx in range(len(boxes2) - 1):
        draw_arrow(draw, (boxes2[idx][0][2], mid_y + 60), (boxes2[idx + 1][0][0], mid_y + 60))

    bottom_y = 570
    boxes3 = [
        ((210, bottom_y, 500, bottom_y + 110), "QA / Repair\n材料与连通检查", "#FFF4DF"),
        ((620, bottom_y, 910, bottom_y + 110), "Optimizer\n命令压缩", "#FFF4DF"),
        ((1030, bottom_y, 1320, bottom_y + 110), "Exporter\nmcfunction", "#FFF4DF"),
        ((1440, bottom_y, 1690, bottom_y + 110), "datapack\n游戏内执行", "#FFF4DF"),
    ]
    for box, text, fill in boxes3:
        draw_round_box(draw, box, text, box_font, fill)
    draw_arrow(draw, (1415, mid_y + 120), (355, bottom_y))
    for idx in range(len(boxes3) - 1):
        draw_arrow(draw, (boxes3[idx][0][2], bottom_y + 55), (boxes3[idx + 1][0][0], bottom_y + 55))

    draw.text((72, 325), "LLM 层：只输出语义 JSON，不直接写坐标", font=small_font, fill="#315B8A")
    draw.text((72, 525), "程序层：确定性生成坐标、房间、路径和命令", font=small_font, fill="#2E7D5B")
    save_diagram(image, "generation_pipeline")


def make_quality_loop():
    image, draw = new_canvas()
    box_font = load_image_font(22, bold=True)
    draw_title(draw, "质量检查与导出闭环", "把一次生成结果变成可复现、可检查、可安装的数据包")

    boxes = [
        ((90, 205, 350, 320), "语义 JSON\n风格 / 拓扑", "#EAF3FF"),
        ((455, 205, 715, 320), "几何网格\n外壳 / 房间 / 路径", "#EAF7EF"),
        ((820, 205, 1080, 320), "QA 检查\n材料 / 入口 / 连通", "#FFF4DF"),
        ((1185, 205, 1445, 320), "Repair\n补洞 / 修路 / 清障", "#FFF4DF"),
        ((1185, 510, 1445, 625), "Optimizer\n命令压缩", "#FFF4DF"),
        ((820, 510, 1080, 625), "Exporter\nmcfunction", "#EAF7EF"),
        ((455, 510, 715, 625), "datapack\n游戏内执行", "#EAF7EF"),
        ((90, 510, 350, 625), "报告与预览\nrun_report / HTML", "#F6F8FA"),
    ]
    for box, text, fill in boxes:
        draw_round_box(draw, box, text, box_font, fill)
    for idx in range(3):
        draw_arrow(draw, (boxes[idx][0][2], 262), (boxes[idx + 1][0][0], 262))
    draw_arrow(draw, (1315, 320), (1315, 510))
    for idx in [4, 5, 6]:
        draw_arrow(draw, (boxes[idx][0][0], 568), (boxes[idx + 1][0][2], 568))
    draw_arrow(draw, (585, 510), (950, 320))
    draw_arrow(draw, (950, 510), (1315, 320))

    note_font = load_image_font(19)
    draw_round_box(
        draw,
        (525, 705, 1275, 790),
        "可复现控制：prompt + seed + blueprint + datapack 一一对应，便于复查和提交",
        note_font,
        "#FFFFFF",
        outline="#B8C3CC",
        radius=16,
    )
    save_diagram(image, "quality_loop")


def make_iteration_timeline():
    image, draw = new_canvas(height=760)
    box_font = load_image_font(20, bold=True)
    small_font = load_image_font(18)
    draw_title(draw, "构思与实现迭代", "从“生成一栋房子”转向“可复现的智能体建筑流水线”")

    y = 360
    milestones = [
        (110, "5/28\n初始 agent\n自动安装与建造"),
        (410, "6/02\n拆分蓝图生成\n子智能体"),
        (710, "6/09\n语义规划层\n多 Agent 架构"),
        (1010, "6/17\nconstruction_method_v1\n本地几何闭环"),
        (1310, "6/18-19\n质量检查\n提交整理"),
    ]
    draw.line((125, y, 1500, y), fill="#345B7D", width=5)
    draw_arrow(draw, (1500, y), (1620, y), width=5)
    for x, label in milestones:
        draw.ellipse((x - 16, y - 16, x + 16, y + 16), fill="#315B8A")
        box = (x - 90, y - 175, x + 150, y - 45) if x % 600 != 410 else (x - 90, y + 45, x + 150, y + 175)
        draw_round_box(draw, box, label, box_font, "#F6F8FA", outline="#9AA9B7", radius=16)
        line_y = box[3] if box[1] < y else box[1]
        draw.line((x, line_y, x, y - 18 if box[1] < y else y + 18), fill="#9AA9B7", width=3)

    draw_round_box(
        draw,
        (390, 610, 1410, 700),
        "核心修改：LLM 只给语义与设计意图；几何、校验、导出由本地确定性模块完成",
        small_font,
        "#EAF7EF",
        outline="#8AA1B7",
        radius=16,
    )
    save_diagram(image, "iteration_timeline")


def make_run_screenshot_collages():
    exterior_items = []
    interior_items = []
    for case in SCREENSHOT_CASES:
        exterior_items.append({"image": f"{case['id']}1", "label": f"{case['id']}1 外景：{case['title']}"})
        exterior_items.append({"image": f"{case['id']}2", "label": f"{case['id']}2 外景：{case['title']}"})
        interior_items.append({"image": f"{case['id']}3", "label": f"{case['id']}3 内饰：{case['title']}"})
    make_screenshot_grid(
        "run_exterior_grid",
        exterior_items,
        "真实运行截图：外景",
        "A-D 对应四个中文 prompt，每组 1/2 为外景截图",
        columns=2,
    )
    make_screenshot_grid(
        "run_interior_grid",
        interior_items,
        "真实运行截图：内饰",
        "A-D 对应四个中文 prompt，每组 3 为室内截图",
        columns=2,
    )


def generate_report_figures():
    make_overall_architecture()
    make_generation_pipeline()
    make_quality_loop()
    make_iteration_timeline()
    make_run_screenshot_collages()


def set_rfonts(run, ascii_name="Times New Roman", east_asia="SimSun"):
    run.font.name = ascii_name
    rfonts = run._element.get_or_add_rPr().rFonts
    rfonts.set(qn("w:ascii"), ascii_name)
    rfonts.set(qn("w:hAnsi"), ascii_name)
    rfonts.set(qn("w:eastAsia"), east_asia)


def set_run_font(run, size=None, bold=None, color=None, ascii_name="Times New Roman", east_asia="SimSun"):
    set_rfonts(run, ascii_name=ascii_name, east_asia=east_asia)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def style_document(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")
    normal.font.size = Pt(11)
    normal.paragraph_format.first_line_indent = Pt(22)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.333
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    for style_name, size, before, after in [
        ("Heading 1", 14, 16, 8),
        ("Heading 2", 12, 10, 5),
        ("Heading 3", 11, 8, 4),
    ]:
        style = doc.styles[style_name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "SimHei")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string("000000")
        style.paragraph_format.first_line_indent = Pt(0)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.25

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.paragraph_format.first_line_indent = Pt(0)
    run = footer.add_run("Minecraft Constructing Agents 课程项目论文")
    set_run_font(run, size=9, color="666666", east_asia="SimSun")


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    tc_pr.append(shading)


def set_cell_width(cell, width_dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:type"), "dxa")
    tc_w.set(qn("w:w"), str(width_dxa))


def set_table_geometry(table, widths_in):
    widths_dxa = [round(width * 1440) for width in widths_in]
    total_dxa = sum(widths_dxa)
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(total_dxa))

    tbl_layout = tbl_pr.first_child_found_in("w:tblLayout")
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")

    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_ind.set(qn("w:w"), "120")

    grid = tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            set_cell_width(cell, widths_dxa[idx])


def add_para(doc, text, style=None, first_line=True, align=None):
    para = doc.add_paragraph(style=style)
    para.alignment = align if align is not None else WD_ALIGN_PARAGRAPH.JUSTIFY
    para.paragraph_format.first_line_indent = Pt(22 if first_line else 0)
    run = para.add_run(text)
    set_run_font(run, size=11)
    return para


def add_centered_line(doc, text, size=10.5, bold=False, after=2):
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.first_line_indent = Pt(0)
    para.paragraph_format.space_after = Pt(after)
    run = para.add_run(text)
    set_run_font(run, size=size, bold=bold, east_asia="SimSun")


def add_docx_table(doc, table_data):
    caption = doc.add_paragraph()
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.first_line_indent = Pt(0)
    caption.paragraph_format.space_before = Pt(4)
    caption.paragraph_format.space_after = Pt(4)
    run = caption.add_run(table_data["caption"])
    set_run_font(run, size=10.5, bold=True, east_asia="SimSun")

    rows = [table_data["headers"]] + table_data["rows"]
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Table Grid"
    set_table_geometry(table, table_data["widths"])
    for r_idx, row in enumerate(rows):
        for c_idx, value in enumerate(row):
            cell = table.cell(r_idx, c_idx)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if r_idx == 0:
                set_cell_shading(cell, "F2F2F2")
            para = cell.paragraphs[0]
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER if c_idx < 2 else WD_ALIGN_PARAGRAPH.LEFT
            para.paragraph_format.first_line_indent = Pt(0)
            para.paragraph_format.space_after = Pt(0)
            run = para.add_run(value)
            set_run_font(run, size=9.5, bold=(r_idx == 0), east_asia="SimSun")
    doc.add_paragraph()


FIGURE_WIDTHS_IN = {
    "overall_architecture": 6.15,
    "generation_pipeline": 5.85,
    "quality_loop": 6.15,
    "iteration_timeline": 5.95,
    "run_exterior_grid": 5.9,
    "run_interior_grid": 6.1,
}


def add_docx_figures(doc, figures):
    for figure in figures:
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para.paragraph_format.first_line_indent = Pt(0)
        run = para.add_run()
        run.add_picture(str(image_path(figure["id"])), width=Inches(FIGURE_WIDTHS_IN.get(figure["id"], 6.1)))
        para.paragraph_format.space_after = Pt(2)

        caption = doc.add_paragraph()
        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
        caption.paragraph_format.first_line_indent = Pt(0)
        caption.paragraph_format.space_before = Pt(0)
        caption.paragraph_format.space_after = Pt(8)
        run = caption.add_run(figure["caption"])
        set_run_font(run, size=10.5, bold=True, east_asia="SimSun")


def build_docx():
    doc = Document()
    style_document(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.first_line_indent = Pt(0)
    title.paragraph_format.space_before = Pt(6)
    title.paragraph_format.space_after = Pt(10)
    run = title.add_run(PAPER["title"])
    set_run_font(run, size=16, bold=True, east_asia="SimHei")

    add_centered_line(doc, PAPER["authors"], size=10.5)
    add_centered_line(doc, PAPER["course"], size=10.5)
    add_centered_line(doc, PAPER["date"], size=10.5, after=10)

    abstract_heading = doc.add_paragraph()
    abstract_heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    abstract_heading.paragraph_format.first_line_indent = Pt(0)
    abstract_heading.paragraph_format.space_after = Pt(3)
    run = abstract_heading.add_run("摘要")
    set_run_font(run, size=11, bold=True, east_asia="SimHei")
    add_para(doc, PAPER["abstract"], first_line=True)
    add_para(doc, f"关键词：{PAPER['keywords']}", first_line=False)

    for section in PAPER["sections"]:
        doc.add_heading(section["heading"], level=1)
        for paragraph in section.get("paragraphs", []):
            add_para(doc, paragraph)
        if section.get("figures"):
            add_docx_figures(doc, section["figures"])
        for subsection in section.get("subsections", []):
            doc.add_heading(subsection["heading"], level=2)
            for paragraph in subsection["paragraphs"]:
                add_para(doc, paragraph)
            if subsection.get("figures"):
                add_docx_figures(doc, subsection["figures"])
        if section.get("table"):
            add_docx_table(doc, section["table"])
        if section.get("figures_after_table"):
            add_docx_figures(doc, section["figures_after_table"])
        if section.get("references"):
            for idx, item in enumerate(section["references"], start=1):
                para = doc.add_paragraph()
                para.paragraph_format.first_line_indent = Pt(0)
                para.paragraph_format.left_indent = Pt(18)
                para.paragraph_format.first_line_indent = Pt(-18)
                para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                run = para.add_run(f"[{idx}] {item}")
                set_run_font(run, size=10.5, east_asia="SimSun")

    OUT_DIR.mkdir(exist_ok=True)
    doc.save(DOCX_PATH)


def register_pdf_fonts():
    candidates = [
        ("CNSong", "CNHei", Path("C:/Windows/Fonts/Deng.ttf"), Path("C:/Windows/Fonts/Dengb.ttf")),
        ("CNSong", "CNHei", Path("C:/Windows/Fonts/msyh.ttc"), Path("C:/Windows/Fonts/msyhbd.ttc")),
        ("CNSong", "CNHei", Path("C:/Windows/Fonts/simsun.ttc"), Path("C:/Windows/Fonts/simhei.ttf")),
        ("CNSong", "CNHei", Path("C:/Windows/Fonts/simsunb.ttf"), Path("C:/Windows/Fonts/simhei.ttf")),
        ("CNSong", "CNHei", Path("C:/Windows/Fonts/simfang.ttf"), Path("C:/Windows/Fonts/simhei.ttf")),
    ]
    for regular_name, bold_name, regular_path, bold_path in candidates:
        if not regular_path.exists() or not bold_path.exists():
            continue
        try:
            pdfmetrics.registerFont(TTFont(regular_name, str(regular_path)))
            pdfmetrics.registerFont(TTFont(bold_name, str(bold_path)))
            return regular_name, bold_name
        except Exception:
            continue
    raise FileNotFoundError("No Chinese-capable Windows font found for PDF generation.")


def pdf_styles(font_name, bold_name):
    styles = getSampleStyleSheet()
    return {
        "title": ParagraphStyle(
            "PaperTitle",
            parent=styles["Title"],
            fontName=bold_name,
            fontSize=16,
            leading=22,
            alignment=TA_CENTER,
            textColor=colors.black,
            spaceAfter=10,
        ),
        "meta": ParagraphStyle(
            "PaperMeta",
            parent=styles["Normal"],
            fontName=font_name,
            fontSize=10.5,
            leading=15,
            alignment=TA_CENTER,
            spaceAfter=2,
        ),
        "abstract_head": ParagraphStyle(
            "PaperAbstractHead",
            parent=styles["Heading2"],
            fontName=bold_name,
            fontSize=11,
            leading=15,
            alignment=TA_LEFT,
            spaceBefore=8,
            spaceAfter=3,
        ),
        "h1": ParagraphStyle(
            "PaperHeading1",
            parent=styles["Heading1"],
            fontName=bold_name,
            fontSize=13,
            leading=18,
            textColor=colors.black,
            spaceBefore=12,
            spaceAfter=6,
            keepWithNext=True,
        ),
        "h2": ParagraphStyle(
            "PaperHeading2",
            parent=styles["Heading2"],
            fontName=bold_name,
            fontSize=11,
            leading=16,
            textColor=colors.black,
            spaceBefore=8,
            spaceAfter=4,
            keepWithNext=True,
        ),
        "body": ParagraphStyle(
            "PaperBody",
            parent=styles["BodyText"],
            fontName=font_name,
            fontSize=10.5,
            leading=16,
            firstLineIndent=21,
            alignment=TA_JUSTIFY,
            spaceAfter=5,
            wordWrap="CJK",
        ),
        "body_no_indent": ParagraphStyle(
            "PaperBodyNoIndent",
            parent=styles["BodyText"],
            fontName=font_name,
            fontSize=10.5,
            leading=16,
            firstLineIndent=0,
            alignment=TA_JUSTIFY,
            spaceAfter=5,
            wordWrap="CJK",
        ),
        "caption": ParagraphStyle(
            "PaperCaption",
            parent=styles["BodyText"],
            fontName=bold_name,
            fontSize=9.5,
            leading=13,
            alignment=TA_CENTER,
            spaceBefore=5,
            spaceAfter=4,
        ),
        "table": ParagraphStyle(
            "PaperTable",
            parent=styles["BodyText"],
            fontName=font_name,
            fontSize=8.6,
            leading=11.5,
            wordWrap="CJK",
        ),
        "table_head": ParagraphStyle(
            "PaperTableHead",
            parent=styles["BodyText"],
            fontName=bold_name,
            fontSize=8.8,
            leading=11.5,
            alignment=TA_CENTER,
        ),
        "ref": ParagraphStyle(
            "PaperReference",
            parent=styles["BodyText"],
            fontName=font_name,
            fontSize=9.5,
            leading=13,
            leftIndent=16,
            firstLineIndent=-16,
            spaceAfter=3,
            wordWrap="CJK",
        ),
    }


def escaped(text):
    return text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")


def P(text, style):
    return Paragraph(escaped(text), style)


def add_pdf_table(story, table_data, styles):
    story.append(P(table_data["caption"], styles["caption"]))
    header = [P(value, styles["table_head"]) for value in table_data["headers"]]
    body = [[P(value, styles["table"]) for value in row] for row in table_data["rows"]]
    widths = [width * inch for width in table_data["widths"]]
    table = Table([header] + body, colWidths=widths, repeatRows=1, hAlign="CENTER")
    table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#F2F2F2")),
                ("GRID", (0, 0), (-1, -1), 0.35, colors.HexColor("#B8B8B8")),
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                ("ALIGN", (0, 0), (1, -1), "CENTER"),
                ("ALIGN", (2, 1), (2, -1), "LEFT"),
                ("TOPPADDING", (0, 0), (-1, -1), 4),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
                ("LEFTPADDING", (0, 0), (-1, -1), 5),
                ("RIGHTPADDING", (0, 0), (-1, -1), 5),
            ]
        )
    )
    story.append(table)
    story.append(Spacer(1, 8))


def add_pdf_figures(story, figures, styles):
    for figure in figures:
        path = image_path(figure["id"])
        with PILImage.open(path) as image:
            aspect = image.height / image.width
        width = FIGURE_WIDTHS_IN.get(figure["id"], 6.1) * inch
        story.append(
            KeepTogether(
                [
                    RLImage(str(path), width=width, height=width * aspect),
                    P(figure["caption"], styles["caption"]),
                    Spacer(1, 8),
                ]
            )
        )


def footer(canvas, doc):
    canvas.saveState()
    canvas.setFont("CNSong", 8)
    canvas.setFillColor(colors.HexColor("#666666"))
    canvas.drawCentredString(letter[0] / 2, 0.5 * inch, f"- {doc.page} -")
    canvas.restoreState()


def build_pdf():
    font_name, bold_name = register_pdf_fonts()
    styles = pdf_styles(font_name, bold_name)
    doc = SimpleDocTemplate(
        str(PDF_PATH),
        pagesize=letter,
        rightMargin=1.0 * inch,
        leftMargin=1.0 * inch,
        topMargin=0.85 * inch,
        bottomMargin=0.8 * inch,
    )

    story = [
        P(PAPER["title"], styles["title"]),
        P(PAPER["authors"], styles["meta"]),
        P(PAPER["course"], styles["meta"]),
        P(PAPER["date"], styles["meta"]),
        Spacer(1, 8),
        P("摘要", styles["abstract_head"]),
        P(PAPER["abstract"], styles["body"]),
        P(f"关键词：{PAPER['keywords']}", styles["body_no_indent"]),
    ]

    for section in PAPER["sections"]:
        story.append(P(section["heading"], styles["h1"]))
        for paragraph in section.get("paragraphs", []):
            story.append(P(paragraph, styles["body"]))
        if section.get("figures"):
            add_pdf_figures(story, section["figures"], styles)
        for subsection in section.get("subsections", []):
            story.append(P(subsection["heading"], styles["h2"]))
            for paragraph in subsection["paragraphs"]:
                story.append(P(paragraph, styles["body"]))
            if subsection.get("figures"):
                add_pdf_figures(story, subsection["figures"], styles)
        if section.get("table"):
            add_pdf_table(story, section["table"], styles)
        if section.get("figures_after_table"):
            add_pdf_figures(story, section["figures_after_table"], styles)
        if section.get("references"):
            for idx, item in enumerate(section["references"], start=1):
                story.append(P(f"[{idx}] {item}", styles["ref"]))

    OUT_DIR.mkdir(exist_ok=True)
    doc.build(story, onFirstPage=footer, onLaterPages=footer)


def main():
    generate_report_figures()
    build_docx()
    build_pdf()
    print(f"DOCX: {DOCX_PATH}")
    print(f"PDF: {PDF_PATH}")


if __name__ == "__main__":
    main()
